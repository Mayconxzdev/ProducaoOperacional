from __future__ import annotations

import json
import sqlite3
import sys
import zipfile
from dataclasses import replace
from datetime import date, timedelta
from pathlib import Path

import pytest
from docx import Document
from PySide6.QtCore import Qt
from PySide6.QtGui import QFontMetrics
from PySide6.QtWidgets import QLabel, QScrollArea, QToolBar

from kanban_app.application.deadline_alert_service import DeadlineAlertService
from kanban_app.application.document_import_service import DocumentImportService
from kanban_app.application.op_discovery_service import OpDiscoveryService
from kanban_app.application.dto import CheckEntryDTO, OpFormDTO, StationRoleDTO
from kanban_app.bootstrap import AppContainer
from kanban_app.domain.enums import CheckState, OpStatus
from kanban_app.infrastructure.config import OpDiscoveryConfig, SmtpConfig, load_config
from kanban_app.infrastructure.db.repositories import ProductionRepository
from kanban_app.infrastructure.db.session import Database
from kanban_app.infrastructure.services.station_runtime import StationRuntimeStore
from kanban_app.main import _application_icon_path, _arguments, _uses_demo_mode
from kanban_app.presentation.widgets.op_form_dialog import OpFormDialog
from kanban_app.presentation.main_window import MainWindow
from kanban_app.presentation.widgets.personalization_dialog import PersonalizationDialog
from kanban_app.presentation.widgets.tv_focus_window import TvFocusWindow
from kanban_app.presentation.tv_settings import default_tv_settings, normalize_tv_settings


def make_container(tmp_path: Path) -> AppContainer:
    tmp_path.mkdir(parents=True, exist_ok=True)
    config = tmp_path / "settings.json"
    config.write_text(
        json.dumps(
            {
                "nas_root": str(tmp_path),
                "database_path": str(tmp_path / "data" / "producao.db"),
                "backups_dir": str(tmp_path / "backups"),
                "state_dir": str(tmp_path / "state"),
                "imports_incoming_dir": str(tmp_path / "imports"),
                "theme_mode": "system",
                "smtp": {"enabled": False},
            }
        ),
        encoding="utf-8",
    )
    return AppContainer.create(config)


def valid_form(container: AppContainer, **changes) -> OpFormDTO:
    sector = container.production_service.sectors(active_only=True)[0]
    base = OpFormDTO(
        numero_op="5059",
        cliente="Cliente",
        modelo="Modelo",
        quantidade=3,
        voltagem="220/380",
        data_inicio=date.today(),
        data_entrega=date.today() + timedelta(days=20),
        setor_id=sector.id,
        status=OpStatus.EM_DIA,
    )
    return replace(base, **changes)


def write_discovery_document(root: Path, group: str, folder: str, *, number: str, delivery: str = "25/12/2026") -> Path:
    target = (
        root
        / "Clientes"
        / "00_PRODUZINDO"
        / group
        / folder
        / "OP"
        / f"OP {number}.docx"
    )
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(
        f"OP: {number}\nCliente: Cliente Automático\nModelo: Ventilador Teste\n"
        f"Quantidade: 2\nVoltagem: 220 V\nPrazo de entrega: {delivery}"
    )
    document.save(target)
    return target


def make_discovery_service(container: AppContainer, source_root: Path) -> OpDiscoveryService:
    config = OpDiscoveryConfig(
        enabled=True,
        source_root_candidates=(source_root,),
        production_relative_path=Path("Clientes/00_PRODUZINDO"),
        groups=("00_GRUPO_A", "00_GRUPO_B"),
        initial_sector_name="Projeto",
        worker_lease_minutes=5,
    )
    return OpDiscoveryService(container.repository, container.document_import_service, config, station_id=container.station_id)


def write_text_pdf(path: Path, text: str) -> None:
    """Cria um PDF textual mínimo para testar a importação sem arquivos externos."""
    lines = [line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)") for line in text.splitlines()]
    commands = [b"BT", b"/F1 12 Tf", b"72 720 Td", b"14 TL"]
    for index, line in enumerate(lines):
        if index:
            commands.append(b"T*")
        commands.append(f"({line}) Tj".encode("latin-1", errors="replace"))
    commands.append(b"ET")
    content = b"\n".join(commands)
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    pdf.extend(b"".join(f"{offset:010d} 00000 n \n".encode("ascii") for offset in offsets[1:]))
    pdf.extend(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii"))
    path.write_bytes(pdf)


def test_demo_uses_an_isolated_local_database_and_seeds_ten_fictitious_ops(tmp_path: Path):
    real_root = tmp_path / "real"
    real_root.mkdir()
    real = make_container(real_root)
    real.production_service.create(valid_form(real, numero_op="800001"))

    demo_root = tmp_path / "demo-local"
    demo = AppContainer.create_demo(demo_root)

    demo_ops = demo.production_service.list_active()
    assert demo.is_demo
    assert demo.config.nas_root == demo_root.resolve()
    assert demo.config.database_path.is_relative_to(demo_root.resolve())
    assert demo.runtime_store.root.is_relative_to(demo_root.resolve())
    assert not demo.config.smtp.enabled
    assert len(demo_ops) == 10
    assert {op.numero_op for op in demo_ops} == {str(number) for number in range(91001, 91011)}
    assert [op.numero_op for op in real.production_service.list_active()] == ["800001"]


def test_demo_reset_restores_exact_seed_without_touching_real_data(tmp_path: Path):
    real_root = tmp_path / "real"
    real_root.mkdir()
    real = make_container(real_root)
    real.production_service.create(valid_form(real, numero_op="800002"))

    demo = AppContainer.create_demo(tmp_path / "demo-local")
    demo.production_service.create(valid_form(demo, numero_op="91999"))
    demo.repository.set_setting("tv.lines_per_page", 3, station_id=demo.station_id)

    restored = demo.reset_demo()

    assert len(restored.production_service.list_active()) == 10
    assert restored.production_service.duplicates("91999") == []
    assert restored.repository.get_setting("tv.lines_per_page") == 10
    assert [op.numero_op for op in real.production_service.list_active()] == ["800002"]


def test_demo_arguments_and_toolbar_expose_safe_guidance(qtbot, tmp_path: Path):
    assert _application_icon_path().is_file()
    assert _arguments(["--demo"]).demo
    assert _arguments(["--configure-role", "demo"]).configure_role == "demo"
    assert _uses_demo_mode(_arguments([]), StationRoleDTO(role="demo"))
    assert _uses_demo_mode(_arguments(["--configure-role", "demo"]), StationRoleDTO(role="office"))
    assert not _uses_demo_mode(_arguments(["--configure-role", "office"]), StationRoleDTO(role="demo"))
    demo = AppContainer.create_demo(tmp_path / "demo-local")
    window = MainWindow(demo)
    qtbot.addWidget(window)
    window._refresh_timer.stop()
    window._deadline_timer.stop()
    toolbar = window.findChild(QToolBar)
    labels = [action.text() for action in toolbar.actions()]
    assert "Guia da demonstração" in labels
    assert "Restaurar 10 OPs fictícias" in labels
    assert window.findChild(QLabel, "demoNotice").text().startswith("MODO DEMONSTRAÇÃO")


def test_application_icon_path_uses_pyinstaller_bundle_assets(monkeypatch, tmp_path: Path):
    bundle_assets = tmp_path / "assets"
    bundle_assets.mkdir()
    expected = bundle_assets / "producao_operacional.png"
    expected.write_bytes(b"icon")
    monkeypatch.setattr(sys, "frozen", True, raising=False)
    monkeypatch.setattr(sys, "_MEIPASS", str(tmp_path), raising=False)

    assert _application_icon_path() == expected


def test_creation_allows_empty_fields_and_status_is_manual(tmp_path: Path):
    container = make_container(tmp_path)
    form = container.production_service.form_defaults()
    created = container.production_service.create(form)
    assert created.numero_op == ""
    assert created.status is OpStatus.EM_DIA
    overdue = container.production_service.update(created.id, replace(form, data_entrega=date.today() - timedelta(days=1), status=OpStatus.AGUARDANDO), created.row_version)
    assert overdue.status is OpStatus.AGUARDANDO


def test_repository_accepts_status_serialized_by_qt(tmp_path: Path):
    container = make_container(tmp_path)
    created = container.production_service.create(replace(valid_form(container), status="EM_DIA"))
    assert created.status is OpStatus.EM_DIA
    updated = container.production_service.update(created.id, replace(valid_form(container), status="AGUARDANDO"), created.row_version)
    assert updated.status is OpStatus.AGUARDANDO


@pytest.mark.parametrize("changes", [{"numero_op": "OP-5"}, {"quantidade": 0}, {"quantidade": -1}])
def test_form_validation_when_values_are_provided(tmp_path: Path, changes):
    container = make_container(tmp_path)
    with pytest.raises(ValueError):
        container.production_service.create(valid_form(container, **changes))


def test_duplicate_number_is_blocked_for_all_new_creations(tmp_path: Path):
    container = make_container(tmp_path)
    container.production_service.create(valid_form(container))
    with pytest.raises(ValueError, match="Já existe uma OP"):
        container.production_service.create(valid_form(container, cliente="Outro cliente"))


def test_check_acompanhamento_has_fourteen_independent_states_and_history(tmp_path: Path):
    container = make_container(tmp_path)
    created = container.production_service.create(valid_form(container, acompanhamento=(CheckEntryDTO("Projetos:IT", CheckState.SIM), CheckEntryDTO("Expedição:Destino", CheckState.NAO))))
    assert len(created.acompanhamento) == 14
    by_key = {entry.field_key: entry for entry in created.acompanhamento}
    assert by_key["Projetos:IT"].state is CheckState.SIM
    assert by_key["Expedição:Destino"].state is CheckState.NAO
    assert all(entry.state in set(CheckState) for entry in created.acompanhamento)
    assert any(event.event_type == "ACOMPANHAMENTO" for event in container.repository.history_for_op(created.id))


def test_completion_reopen_archive_restore_and_history(tmp_path: Path):
    container = make_container(tmp_path)
    created = container.production_service.create(valid_form(container))
    completed = container.production_service.complete(created.id, created.row_version)
    assert completed.status is OpStatus.CONCLUIDO
    assert not container.production_service.list_active()
    reopened = container.production_service.reopen(completed.id, completed.row_version, OpStatus.PRIORIDADE)
    assert reopened.status is OpStatus.PRIORIDADE and reopened.completed_at is None
    archived = container.production_service.archive(reopened.id, reopened.row_version)
    assert container.repository.list_archived_ops()[0].id == archived.id
    restored = container.production_service.restore(archived.id, archived.row_version)
    assert restored.archived is False
    assert any(event.event_type == "ARQUIVADA" for event in container.repository.history_for_op(created.id))


def test_sectors_are_shared_renamable_and_delete_migrates_ops(tmp_path: Path):
    container = make_container(tmp_path)
    sectors = container.production_service.sectors(active_only=True)
    created = container.production_service.create(valid_form(container, setor_id=sectors[0].id))
    container.repository.update_sector(sectors[0].id, nome="Projetos", cor="#123456", ativo=True, ordem=1, station_id="test")
    assert container.production_service.get(created.id).setor_nome == "Projetos"
    container.repository.delete_sector_with_migration(sectors[0].id, sectors[1].id, station_id="test")
    assert container.production_service.get(created.id).setor_id == sectors[1].id


def test_deadline_colors_are_independent_of_status(tmp_path: Path):
    container = make_container(tmp_path)
    service = container.production_service
    sector = service.sectors(active_only=True)[0]
    assert service.deadline_band(service.create(valid_form(container, setor_id=sector.id, data_entrega=date.today() + timedelta(days=15), status=OpStatus.PRIORIDADE))) == "sector"
    warning = service.create(valid_form(container, numero_op="5060", data_entrega=date.today() + timedelta(days=8), status=OpStatus.CONCLUIDO))
    critical = service.create(valid_form(container, numero_op="5061", data_entrega=date.today() - timedelta(days=1), status=OpStatus.EM_DIA))
    assert service.deadline_band(warning) == "warning"
    assert service.deadline_band(critical) == "critical"


def test_document_import_docx_odt_and_pdf_sample(tmp_path: Path):
    container = make_container(tmp_path)
    importer = DocumentImportService()
    sector_id = container.production_service.sectors(active_only=True)[0].id
    docx_path = tmp_path / "op.docx"
    document = Document()
    document.add_paragraph("OP: 5209\nCliente: ACME Industrial\nModelo: AEROFLOW PE 200e T4 0,25 220/380 V\nQuantidade: 10\nPrazo de entrega: 25/04/2026")
    document.save(docx_path)
    imported_docx = importer.extract(docx_path, default_sector_id=sector_id)
    assert imported_docx.form.numero_op == "5209"
    assert imported_docx.form.quantidade == 10
    assert imported_docx.form.voltagem == "220/380"
    assert "220/380 V" not in imported_docx.form.modelo
    odt_path = tmp_path / "op.odt"
    with zipfile.ZipFile(odt_path, "w") as archive:
        archive.writestr("content.xml", "<document><p>OP: 5210</p><p>Cliente: Beta Equipamentos</p><p>Quantidade: 1</p><p>Prazo: 01/05/2026</p></document>")
    imported_odt = importer.extract(odt_path, default_sector_id=sector_id)
    assert imported_odt.form.numero_op == "5210"
    pdf_path = tmp_path / "op-5292.pdf"
    write_text_pdf(
        pdf_path,
        "OP: 5292\nCLIENTE: Alfa Equipamentos\nMODELO: AEROFLOW PE 200e T4 0,25 220 V\nQUANTIDADE: 1\nPRAZO DE ENTREGA: 03/07/2026",
    )
    imported_pdf = importer.extract(pdf_path, default_sector_id=sector_id)
    assert not imported_pdf.errors
    assert imported_pdf.form.numero_op == "5292"
    assert imported_pdf.form.cliente == "Alfa Equipamentos"
    assert imported_pdf.form.modelo == "AEROFLOW PE 200e T4 0,25"
    assert imported_pdf.form.quantidade == 1
    assert imported_pdf.form.voltagem == "220"
    assert imported_pdf.form.data_entrega == date(2026, 7, 3)


def test_odt_fragmented_spans_preserve_op_fields(tmp_path: Path):
    importer = DocumentImportService()
    path = tmp_path / "OP 5299.odt"
    content = """<?xml version="1.0" encoding="UTF-8"?>
    <office:document-content
      xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
      xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0">
      <office:body><office:text>
        <text:p>OP - <text:span>52</text:span><text:span>99</text:span></text:p>
        <text:p>CLIENTE: Nova Era Manutencao Industrial Ltda</text:p>
        <text:p>MODELO: AEROFLOW PE 250e T2 0,33 CV - Trifásico</text:p>
        <text:p>C/ PROTETOR DE CHUVA</text:p>
        <text:p>QUANTIDADE: 6</text:p>
        <text:p>VOLTAGEM: <text:span>22</text:span><text:span>0V</text:span></text:p>
        <text:p>OBS: PRAZO DE ENTREGA <text:span>17/</text:span><text:span>08/2026</text:span></text:p>
      </office:text></office:body>
    </office:document-content>"""
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("content.xml", content)
    preview = importer.extract(path, default_sector_id="setor")
    assert preview.form.numero_op == "5299"
    assert preview.form.cliente == "Nova Era Manutencao Industrial Ltda"
    assert preview.form.modelo == "AEROFLOW PE 250e T2 0,33 CV - Trifásico C/ PROTETOR DE CHUVA"
    assert preview.form.quantidade == 6
    assert preview.form.voltagem == "220"
    assert preview.form.data_entrega == date(2026, 8, 17)
    assert not preview.missing_fields


def test_discovery_first_run_creates_a_baseline_without_reading_or_importing_existing_ops(tmp_path: Path):
    container = make_container(tmp_path / "app")
    root = tmp_path / "nas"
    write_discovery_document(root, "00_GRUPO_A", "01 - OP 7001 - Já Existente", number="7001")
    (root / "Clientes" / "00_PRODUZINDO" / "00_GRUPO_B").mkdir(parents=True)

    service = make_discovery_service(container, root)
    first = service.run()

    assert first.status == "BASELINED"
    assert first.baselined == 1
    assert container.production_service.duplicates("7001") == []

    write_discovery_document(root, "00_GRUPO_A", "02 - OP 7002 - Nova", number="7002")
    second = service.run()

    assert second.status == "OK"
    assert second.imported == 1
    imported = container.production_service.duplicates("7002")
    assert len(imported) == 1
    assert imported[0].setor_nome == "Projeto"
    assert imported[0].status is OpStatus.EM_DIA


def test_discovery_blocks_an_existing_op_number_and_keeps_its_source_record(tmp_path: Path):
    container = make_container(tmp_path / "app")
    root = tmp_path / "nas"
    (root / "Clientes" / "00_PRODUZINDO" / "00_GRUPO_A").mkdir(parents=True)
    (root / "Clientes" / "00_PRODUZINDO" / "00_GRUPO_B").mkdir(parents=True)
    service = make_discovery_service(container, root)
    assert service.run().status == "BASELINED"

    container.production_service.create(valid_form(container, numero_op="7003"))
    source = write_discovery_document(root, "00_GRUPO_A", "03 - OP 7003 - Repetida", number="7003")
    result = service.run()

    assert result.imported == 0
    assert result.blocked_duplicates == 1
    records = container.repository.import_source_records([source.relative_to(root / "Clientes" / "00_PRODUZINDO").as_posix()])
    assert next(iter(records.values())).state == "BLOCKED_DUPLICATE"


def test_discovery_waits_for_required_fields_and_retries_after_document_changes(tmp_path: Path):
    container = make_container(tmp_path / "app")
    root = tmp_path / "nas"
    (root / "Clientes" / "00_PRODUZINDO" / "00_GRUPO_A").mkdir(parents=True)
    (root / "Clientes" / "00_PRODUZINDO" / "00_GRUPO_B").mkdir(parents=True)
    service = make_discovery_service(container, root)
    assert service.run().status == "BASELINED"

    source = write_discovery_document(root, "00_GRUPO_A", "04 - OP 7004 - Incompleta", number="7004", delivery="")
    first = service.run()
    assert first.pending == 1
    assert container.production_service.duplicates("7004") == []

    document = Document(source)
    document.paragraphs[0].add_run("\nPrazo de entrega: 26/12/2026")
    document.save(source)
    second = service.run()
    assert second.imported == 1
    assert len(container.production_service.duplicates("7004")) == 1


def test_flexible_date_field_accepts_digits_and_normalizes(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    dialog = OpFormDialog(sectors=container.production_service.sectors(active_only=True), voltages=[], initial=None)
    qtbot.addWidget(dialog)
    dialog.delivery_date.line_edit.setText("13082026")
    dialog.delivery_date.normalize()
    assert dialog.delivery_date.line_edit.text() == "13/08/2026"
    assert dialog.delivery_date.date_value() == date(2026, 8, 13)
    dialog.start_date.line_edit.setText("13/08/2026")
    dialog.start_date.normalize()
    assert dialog.start_date.date_value() == date(2026, 8, 13)


def test_sector_background_and_text_colors_are_persisted_and_rendered(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    sector = container.production_service.sectors(active_only=True)[0]
    updated_sector = container.repository.update_sector(
        sector.id,
        nome=sector.nome,
        cor="#fef08a",
        cor_texto="#7f1d1d",
        ativo=True,
        ordem=sector.ordem,
        station_id="test",
    )
    assert updated_sector.cor == "#fef08a"
    assert updated_sector.cor_texto == "#7f1d1d"
    op = container.production_service.create(
        valid_form(container, setor_id=sector.id, data_entrega=date.today() + timedelta(days=30))
    )
    tv = TvFocusWindow(settings={"visible_columns": ["op", "setor"]})
    qtbot.addWidget(tv)
    tv.set_ops([op])
    foreground = tv.list_view.model.data(
        tv.list_view.model.index(0, 0), Qt.ItemDataRole.ForegroundRole
    )
    background = tv.list_view.model.data(
        tv.list_view.model.index(0, 0), Qt.ItemDataRole.BackgroundRole
    )
    assert background.color().name() == "#fef08a"
    assert foreground.color().name() == "#7f1d1d"
    tv._timer.stop()


def test_deadline_email_is_consolidated_and_idempotent(tmp_path: Path, monkeypatch):
    container = make_container(tmp_path)
    service = container.production_service
    service.create(valid_form(container, data_entrega=date.today() + timedelta(days=14)))
    service.create(valid_form(container, numero_op="5060", data_entrega=date.today() + timedelta(days=7)))
    container.repository.set_setting("deadline.email_recipients", ["ops@example.com"], station_id="test")
    sent: list[object] = []
    class FakeSmtp:
        def __init__(self, *args, **kwargs): pass
        def __enter__(self): return self
        def __exit__(self, *args): return False
        def login(self, *args): pass
        def send_message(self, message): sent.append(message)
    monkeypatch.setattr("smtplib.SMTP_SSL", FakeSmtp)
    alerts = DeadlineAlertService(container.repository, SmtpConfig(enabled=True, host="smtp.example", from_email="from@example.com"), station_id="test")
    assert alerts.run_daily() == 2
    assert len(sent) == 1
    assert alerts.run_daily() == 0


def test_legacy_migration_creates_backup_preserves_data_and_removes_obsolete_tables(tmp_path: Path):
    db_path = tmp_path / "legacy.db"
    with sqlite3.connect(db_path) as connection:
        connection.executescript(
            "CREATE TABLE app_meta(meta_key TEXT PRIMARY KEY, meta_value TEXT);"
            "CREATE TABLE ops(id INTEGER PRIMARY KEY, empresa_grupo TEXT, numero_op TEXT, cliente TEXT, modelo TEXT, quantidade INTEGER, data_inicio TEXT, data_entrega TEXT, status_geral TEXT, percentual_pronto REAL, pendencia_principal TEXT, observacao_geral TEXT, created_at TEXT, updated_at TEXT, ordem_lista INTEGER, tensao TEXT, setor TEXT, prioridade INTEGER, archived INTEGER, archived_at TEXT, archived_by TEXT, archived_reason TEXT, row_version INTEGER, status_since TEXT);"
            "CREATE TABLE historico(id INTEGER PRIMARY KEY, entidade TEXT, entidade_id INTEGER, campo TEXT, valor_anterior TEXT, valor_novo TEXT, usuario TEXT, data_hora TEXT);"
            "CREATE TABLE app_settings(setting_key TEXT PRIMARY KEY, setting_value TEXT, value_type TEXT, version INTEGER, updated_by TEXT, updated_at TEXT, scope TEXT);"
            "CREATE TABLE app_users(id INTEGER PRIMARY KEY, username TEXT);"
            "CREATE TABLE op_items(id INTEGER PRIMARY KEY, op_id INTEGER);"
            "INSERT INTO ops VALUES(1,'g','5059','Cliente','Modelo',1,'2026-01-01','2026-01-10','RESOLVIDA',0,'Sem pendências','', '2026-01-01','2026-01-02',1,'220V','Projeto',0,0,NULL,NULL,'',1,'2026-01-01');"
            "INSERT INTO historico VALUES(1,'OP',1,'created','','OP criada','admin','2026-01-01');"
        )
    database = Database(db_path, backups_dir=tmp_path / "backups")
    database.create_schema()
    repository = ProductionRepository(database)
    assert repository.list_concluded_ops()[0].status is OpStatus.CONCLUIDO
    assert repository.history_for_op(1)
    assert list((tmp_path / "backups").glob("kanban_pre_v2_*.db"))
    with sqlite3.connect(db_path) as connection:
        tables = {row[0] for row in connection.execute("SELECT name FROM sqlite_master WHERE type='table'")}
    assert "app_users" not in tables and "op_items" not in tables


def test_form_has_two_tabs_and_fourteen_check_controls(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    dialog = OpFormDialog(sectors=container.production_service.sectors(active_only=True), voltages=[], initial=None)
    qtbot.addWidget(dialog)
    dialog.resize(744, 718)
    dialog.show()
    dialog.tabs.setCurrentIndex(1)
    qtbot.wait(30)
    assert dialog.windowTitle() == "Cadastrar OP"
    assert len(dialog._check_inputs) == 14
    assert len(dialog._check_groups) == 3
    assert isinstance(dialog._check_scroll, QScrollArea)
    assert dialog._check_scroll.horizontalScrollBar().maximum() == 0
    assert all(combo.minimumWidth() >= 170 for combo in dialog._check_inputs.values())
    assert all(group.height() > 0 for group in dialog._check_groups)




def test_two_stations_share_ops_and_refresh_interval_is_1_5_seconds(qtbot, tmp_path: Path):
    first = make_container(tmp_path)
    second = AppContainer.create(tmp_path / "settings.json")
    created = first.production_service.create(valid_form(first, numero_op="7001", cliente="Sincronização"))
    assert second.production_service.get(created.id).cliente == "Sincronização"
    window = MainWindow(second)
    qtbot.addWidget(window)
    assert window._refresh_timer.interval() == 1500
    window._refresh_timer.stop()

def test_office_toolbar_and_tv_do_not_expose_removed_controls(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    window = MainWindow(container)
    qtbot.addWidget(window)
    window._refresh_timer.stop()
    labels = [action.text() for action in window.findChild(QToolBar).actions()]
    for required in ("Nova OP", "Importar OP", "Histórico", "Personalização", "Abrir modo TV/Foco"):
        assert required in labels
    assert not {"Ver Kanban", "Aplicar filtros", "Atualizar quadro", "Sair"} & set(labels)
    tv = TvFocusWindow(visible_columns=["op", "status"], page_interval_seconds=13, lines_per_page=8)
    qtbot.addWidget(tv)
    pendencia_index = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "pendencia")
    assert tv.list_view.table.isColumnHidden(pendencia_index)
    tv._timer.stop()


def test_tv_uses_shared_widths_and_fills_height_with_fewer_rows(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    ops = [container.production_service.create(valid_form(container, numero_op=str(5300 + index))) for index in range(3)]
    tv = TvFocusWindow(
        settings={
            "visible_columns": ["op", "modelo", "setor"],
            "column_order": ["modelo", "op", "setor"],
            "column_widths": {"op": 88, "modelo": 420, "setor": 155},
            "lines_per_page": 10,
            "font_scale_percent": 100,
        }
    )
    qtbot.addWidget(tv)
    tv.resize(1000, 700)
    tv.show()
    tv.set_ops(ops)
    qtbot.wait(50)
    op_index = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "op")
    model_index = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "modelo")
    sector_index = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "setor")
    fitted_total = sum(tv.list_view.table.columnWidth(index) for index in (op_index, model_index, sector_index))
    assert abs(fitted_total - tv.list_view.table.viewport().width()) <= 2
    assert tv.list_view.table.columnWidth(model_index) / tv.list_view.table.columnWidth(op_index) == pytest.approx(420 / 88, rel=0.08)
    assert tv.list_view.table.verticalHeader().defaultSectionSize() > 150
    tv._timer.stop()


def test_demo_tv_preset_keeps_visible_values_complete_on_full_hd(qtbot, tmp_path: Path):
    """A vitrine padrão da TV não deve renderizar campos com reticências."""

    demo = AppContainer.create_demo(tmp_path / "demo-tv")
    settings = {
        key: demo.repository.get_setting(f"tv.{key}", value)
        for key, value in default_tv_settings().items()
    }
    tv = TvFocusWindow(settings=settings)
    qtbot.addWidget(tv)
    tv.resize(1920, 1080)
    tv.show()
    tv.set_ops(demo.production_service.list_active())
    qtbot.wait(80)

    table = tv.list_view.table
    padding = int(settings["cell_padding_px"])
    visible_columns = set(settings["visible_columns"])
    overflowing: list[tuple[int, str, str]] = []
    for row in range(table.model().rowCount()):
        for column, (key, _label) in enumerate(tv.list_view.model.COLUMNS):
            if key not in visible_columns:
                continue
            index = table.model().index(row, column)
            text = str(index.data(Qt.ItemDataRole.DisplayRole) or "")
            font = index.data(Qt.ItemDataRole.FontRole)
            available_width = table.visualRect(index).width() - (2 * padding)
            if text and QFontMetrics(font).horizontalAdvance(text) > available_width:
                overflowing.append((row, key, text))

    tv._timer.stop()
    assert not overflowing, f"A TV cortaria valores visíveis: {overflowing}"


def test_personalization_persists_complete_tv_settings_for_other_stations(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    op = container.production_service.create(valid_form(container))
    dialog = PersonalizationDialog(
        container.repository,
        container.config,
        container.station_id,
        preview_ops=[op],
        tv_settings={"lines_per_page": 8},
    )
    qtbot.addWidget(dialog)
    dialog.tv_lines.setValue(10)
    dialog.tv_font.setValue(115)
    op_row = next(row for row in range(dialog.tv_columns.rowCount()) if dialog.tv_columns.item(row, 0).data(Qt.ItemDataRole.UserRole) == "op")
    dialog.tv_columns.cellWidget(op_row, 3).setValue(92)
    voltage_row = next(row for row in range(dialog.tv_columns.rowCount()) if dialog.tv_columns.item(row, 0).data(Qt.ItemDataRole.UserRole) == "voltagem")
    dialog.tv_columns.cellWidget(voltage_row, 2).setText("Volt")
    dialog._save()
    assert container.repository.get_setting("tv.lines_per_page") == 10
    assert container.repository.get_setting("tv.font_scale_percent") == 115
    assert container.repository.get_setting("tv.column_widths")["op"] == 92
    assert container.repository.get_setting("tv.column_headers")["voltagem"] == "Volt"
    assert container.repository.get_setting("tv.column_formats")["entrega"] == "dd/MM/yy"
    dialog.tv_preview._timer.stop()


def test_personalization_saves_op_discovery_monitoring_rules_locally(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    dialog = PersonalizationDialog(container.repository, container.config, container.station_id)
    qtbot.addWidget(dialog)

    assert "Integração de OPs" in [dialog.tabs.tabText(index) for index in range(dialog.tabs.count())]
    dialog.op_discovery_enabled.setChecked(True)
    dialog.op_discovery_roots.setPlainText("\\\\SERVIDOR\\Compartilhamento\nZ:\\")
    dialog.op_discovery_relative.setText("CERTIFICADOS/Clientes/00_PRODUZINDO")
    dialog.op_discovery_groups.setPlainText("00_GRUPO_A\n00_GRUPO_B")
    dialog.op_discovery_pdf.setChecked(False)
    dialog.op_discovery_days["saturday"].setChecked(True)
    dialog.op_discovery_days["monday"].setChecked(False)
    dialog.op_discovery_times.setPlainText("07:30\n16:45")
    dialog._save()

    saved = load_config(tmp_path / "settings.json").op_discovery
    assert saved.enabled
    assert [str(item).rstrip("\\") for item in saved.source_root_candidates] == ["\\\\SERVIDOR\\Compartilhamento", "Z:"]
    assert saved.production_relative_path == Path("CERTIFICADOS/Clientes/00_PRODUZINDO")
    assert saved.groups == ("00_GRUPO_A", "00_GRUPO_B")
    assert saved.document_extensions == (".odt", ".docx")
    assert saved.schedule_days == ("tuesday", "wednesday", "thursday", "friday", "saturday")
    assert saved.schedule_times == ("07:30", "16:45")
    dialog.tv_preview._timer.stop()


def test_personalization_sector_cards_are_colored_without_hover_and_theme_is_local(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    sector = container.production_service.sectors(active_only=True)[0]
    container.repository.update_sector(
        sector.id,
        nome=sector.nome,
        cor="#124e78",
        cor_texto="#fef3c7",
        ativo=True,
        ordem=sector.ordem,
        station_id="test",
    )
    runtime_store = StationRuntimeStore(tmp_path / "runtime")
    dialog = PersonalizationDialog(
        container.repository,
        container.config,
        container.station_id,
        runtime_store=runtime_store,
        theme_mode="system",
    )
    qtbot.addWidget(dialog)
    dialog.show()
    qtbot.wait(30)

    card = dialog.sector_list.itemWidget(dialog.sector_list.item(0))
    assert card.name_label.text() == f"{sector.ordem}. {sector.nome}"
    assert "#124e78" in card.styleSheet().casefold()
    assert "#fef3c7" in card.styleSheet().casefold()
    assert card._selected

    theme_events: list[str] = []
    dialog.office_theme_changed.connect(theme_events.append)
    dialog.office_theme.setCurrentIndex(dialog.office_theme.findData("dark"))
    dialog._save()

    assert runtime_store.load_theme_mode() == "dark"
    assert theme_events == ["dark"]
    dialog.tv_preview._timer.stop()


def test_station_theme_mode_is_validated_and_persisted_locally(tmp_path: Path):
    runtime_store = StationRuntimeStore(tmp_path / "runtime")
    assert runtime_store.load_theme_mode("light") == "light"
    runtime_store.save_theme_mode("dark")
    assert runtime_store.load_theme_mode() == "dark"
    runtime_store.save_theme_mode("not-a-theme")
    assert runtime_store.load_theme_mode("light") == "system"


def test_tv_keeps_pagination_timer_running_when_settings_are_unchanged(qtbot):
    tv = TvFocusWindow(settings={"page_interval_seconds": 2, "lines_per_page": 2})
    qtbot.addWidget(tv)
    qtbot.wait(120)
    before = tv._timer.remainingTime()
    tv.apply_settings(tv.settings)
    after = tv._timer.remainingTime()
    assert tv._timer.isActive()
    assert after <= before + 40
    tv._timer.stop()


def test_tv_uses_deadline_override_and_individual_column_font_scale(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    sector = container.production_service.sectors(active_only=True)[0]
    container.repository.update_sector(sector.id, nome=sector.nome, cor="#12a4c7", ativo=True, ordem=sector.ordem, station_id="test")
    op = container.production_service.create(
        valid_form(container, setor_id=sector.id, data_entrega=date.today() - timedelta(days=2), status=OpStatus.EM_ATRASO)
    )
    tv = TvFocusWindow(
        settings={
            "visible_columns": ["op", "modelo", "setor"],
            "column_font_scales": {"op": 160, "modelo": 75, "setor": 100},
        }
    )
    qtbot.addWidget(tv)
    tv.set_ops([op])
    tv.resize(900, 500)
    tv.show()
    qtbot.wait(30)
    op_column = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "op")
    model_column = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "modelo")
    background = tv.list_view.model.data(tv.list_view.model.index(0, op_column), Qt.ItemDataRole.BackgroundRole)
    op_font = tv.list_view.model.data(tv.list_view.model.index(0, op_column), Qt.ItemDataRole.FontRole)
    model_font = tv.list_view.model.data(tv.list_view.model.index(0, model_column), Qt.ItemDataRole.FontRole)
    assert background.color().name() == "#ef4444"
    assert op_font.pointSize() > model_font.pointSize()
    tv._timer.stop()


def test_import_understands_portuguese_date_quantity_unit_and_filename_number(tmp_path: Path):
    importer = DocumentImportService()
    path = tmp_path / "OP 6401.docx"
    document = Document()
    document.add_paragraph(
        "Cliente: Indústria Exemplo\n"
        "Produto: AEROFLOW PE 500 T4 220-380 VAC\n"
        "Qtd.: 12 peças\n"
        "Previsão de entrega: 17 de agosto de 2026"
    )
    document.save(path)
    preview = importer.extract(path, default_sector_id="setor")
    assert preview.form.numero_op == "6401"
    assert preview.form.quantidade == 12
    assert preview.form.voltagem == "220/380"
    assert preview.form.data_entrega == date(2026, 8, 17)
    assert not preview.missing_fields


def test_import_batch_only_enables_complete_items_after_review(qtbot):
    from kanban_app.application.dto import ImportPreviewDTO
    from kanban_app.presentation.widgets.import_batch_dialog import ImportBatchDialog

    incomplete_form = OpFormDTO(numero_op="7001", cliente="Cliente")
    preview = ImportPreviewDTO("op.docx", incomplete_form, missing_fields=("Modelo", "Quantidade", "Voltagem", "Prazo de entrega"))
    dialog = ImportBatchDialog([preview])
    qtbot.addWidget(dialog)
    assert not dialog.confirm_button.isEnabled()
    complete_form = replace(
        incomplete_form,
        modelo="Modelo",
        quantidade=1,
        voltagem="220",
        data_entrega=date(2026, 9, 1),
    )
    dialog.update_preview(0, replace(preview, form=complete_form, missing_fields=()))
    assert dialog.confirm_button.isEnabled()
    assert dialog.complete_previews()[0].form.modelo == "Modelo"


def test_repository_reads_and_writes_tv_settings_as_one_group(tmp_path: Path):
    container = make_container(tmp_path)
    values = {
        "tv.lines_per_page": 10,
        "tv.visible_columns": ["op", "modelo", "setor"],
        "tv.column_widths": {"op": 90, "modelo": 500, "setor": 150},
    }
    container.repository.set_settings(values, station_id="tv-test")
    assert container.repository.get_settings(tuple(values), defaults={}) == values


def test_tv_allows_individual_header_date_and_sector_text(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    sector = container.production_service.sectors(active_only=True)[0]
    delivery = date(2026, 8, 9)
    op = container.production_service.create(valid_form(container, setor_id=sector.id, data_entrega=delivery))
    tv = TvFocusWindow(
        settings={
            "visible_columns": ["op", "voltagem", "entrega", "setor"],
            "column_headers": {"voltagem": "Volt", "entrega": "Entrega", "setor": "Setor"},
            "column_formats": {"entrega": "dd/MM", "inicio": "dd/MM/yyyy"},
            "sector_labels": {sector.id: "Proj."},
        }
    )
    qtbot.addWidget(tv)
    tv.set_ops([op])
    voltage_column = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "voltagem")
    delivery_column = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "entrega")
    sector_column = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "setor")
    assert tv.list_view.model.headerData(voltage_column, Qt.Orientation.Horizontal, Qt.ItemDataRole.DisplayRole) == "Volt"
    assert tv.list_view.model.headerData(delivery_column, Qt.Orientation.Horizontal, Qt.ItemDataRole.DisplayRole) == "Entrega"
    assert tv.list_view.model.data(tv.list_view.model.index(0, delivery_column), Qt.ItemDataRole.DisplayRole) == "09/08"
    assert tv.list_view.model.data(tv.list_view.model.index(0, sector_column), Qt.ItemDataRole.DisplayRole) == "Proj."
    tv._timer.stop()


def test_tv_sector_filter_selected_is_applied(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    sectors = container.production_service.sectors(active_only=True)
    first = container.production_service.create(valid_form(container, numero_op="6001", setor_id=sectors[0].id))
    second = container.production_service.create(valid_form(container, numero_op="6002", setor_id=sectors[1].id))
    tv = TvFocusWindow(settings={"sector_filter_mode": "selected", "visible_sector_ids": [sectors[1].id]})
    qtbot.addWidget(tv)
    tv.set_ops([first, second])
    assert [op.numero_op for op in tv._display_ops] == ["6002"]
    tv._timer.stop()


def test_tv_settings_migrate_global_abbreviation_to_individual_values():
    settings = normalize_tv_settings({"abbreviate_headers_and_dates": True})
    assert settings["column_headers"]["voltagem"] == "V"
    assert settings["column_headers"]["cliente"] == "Cliente"
    assert settings["column_formats"]["entrega"] == "dd/MM"
    customized = normalize_tv_settings({"column_headers": {"voltagem": "Volt"}})
    assert customized["column_headers"]["voltagem"] == "Volt"
    assert customized["column_headers"]["entrega"] == "Entrega"


def test_tv_delegate_really_paints_deadline_background(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    op = container.production_service.create(
        valid_form(container, data_entrega=date.today() - timedelta(days=1))
    )
    tv = TvFocusWindow(settings={"visible_columns": ["op"], "lines_per_page": 1})
    qtbot.addWidget(tv)
    tv.set_deadline_colors("#f9a8d4", "#ef4444")
    tv.set_deadline_rules(warning_days=14, critical_days=7, eligible_sector_ids=None)
    tv.set_ops([op])
    tv.resize(500, 300)
    tv.show()
    qtbot.wait(40)
    index = tv.list_view.model.index(0, 0)
    rect = tv.list_view.table.visualRect(index)
    image = tv.list_view.table.viewport().grab().toImage()
    color = image.pixelColor(min(rect.right() - 4, rect.left() + 8), min(rect.bottom() - 4, rect.top() + 8))
    assert color.red() > 220 and color.green() < 110 and color.blue() < 110
    tv._timer.stop()


def test_pdf_text_reader_uses_only_first_page(monkeypatch, tmp_path: Path):
    class FakePage:
        def __init__(self, text: str):
            self.text = text

        def extract_text(self):
            return self.text

    class FakeReader:
        def __init__(self, _path: str):
            self.pages = [FakePage("OP: 7001\nCliente: PRIMEIRA"), FakePage("OP: 9999\nCliente: ANEXO")]

    path = tmp_path / "duas_paginas.pdf"
    path.write_bytes(b"fake")
    monkeypatch.setattr("kanban_app.application.document_import_service.PdfReader", FakeReader)
    text = DocumentImportService()._read_text(path)
    assert "7001" in text and "PRIMEIRA" in text
    assert "9999" not in text and "ANEXO" not in text


def test_tv_status_formatting_accepts_legacy_string(qtbot, tmp_path: Path):
    container = make_container(tmp_path)
    op = container.production_service.create(valid_form(container))
    legacy = replace(op, status="EM_DIA")
    tv = TvFocusWindow(settings={"visible_columns": ["op", "status"]})
    qtbot.addWidget(tv)
    tv.set_ops([legacy])
    status_column = next(index for index, column in enumerate(tv.list_view.model.COLUMNS) if column[0] == "status")
    assert tv.list_view.model.data(tv.list_view.model.index(0, status_column), Qt.ItemDataRole.DisplayRole) == "Em dia"
    tv._timer.stop()


def test_schema_24_adds_sector_text_color_and_normalizes_existing_voltage(tmp_path: Path):
    container = make_container(tmp_path)
    sector = container.production_service.sectors(active_only=True)[0]
    created = container.production_service.create(valid_form(container, numero_op="8100", voltagem="440 V"))
    db_path = container.database.db_path
    with sqlite3.connect(db_path) as connection:
        # Simula a estrutura 2.3.4, anterior à cor de texto por setor.
        connection.execute("ALTER TABLE sectors RENAME TO sectors_with_text")
        connection.execute(
            "CREATE TABLE sectors (id VARCHAR(36) PRIMARY KEY, nome VARCHAR(120) NOT NULL, "
            "ordem INTEGER NOT NULL, cor VARCHAR(16) NOT NULL, ativo BOOLEAN NOT NULL, "
            "created_at DATETIME NOT NULL, updated_at DATETIME NOT NULL, UNIQUE(nome))"
        )
        connection.execute(
            "INSERT INTO sectors(id,nome,ordem,cor,ativo,created_at,updated_at) "
            "SELECT id,nome,ordem,cor,ativo,created_at,updated_at FROM sectors_with_text"
        )
        connection.execute("DROP TABLE sectors_with_text")
        connection.execute("UPDATE ops SET voltagem='440 V' WHERE id=?", (created.id,))
        connection.execute("UPDATE app_meta SET meta_value='24' WHERE meta_key='schema_version'")
        connection.commit()
    upgraded = AppContainer.create(tmp_path / "settings.json")
    assert upgraded.database.schema_version() == Database.SCHEMA_VERSION
    restored_sector = next(item for item in upgraded.production_service.sectors() if item.id == sector.id)
    assert restored_sector.cor_texto.startswith("#")
    assert upgraded.production_service.get(created.id).voltagem == "440"


def test_schema_23_with_obsolete_mutation_table_upgrades_and_status_save_works(tmp_path: Path):
    container = make_container(tmp_path)
    created = container.production_service.create(valid_form(container, numero_op="8101"))
    db_path = container.database.db_path
    with sqlite3.connect(db_path) as connection:
        connection.execute(
            "UPDATE app_meta SET meta_value = '23' WHERE meta_key = 'schema_version'"
        )
        connection.execute(
            "CREATE TABLE IF NOT EXISTS op_mutations("
            "id INTEGER PRIMARY KEY, op_id INTEGER NOT NULL, payload_json TEXT NOT NULL)"
        )
        connection.execute(
            "CREATE TABLE IF NOT EXISTS station_presence("
            "machine_id TEXT PRIMARY KEY, last_seen_at TEXT NOT NULL)"
        )
        connection.commit()

    upgraded = AppContainer.create(tmp_path / "settings.json")
    current = upgraded.production_service.get(created.id)
    assert current is not None
    updated = upgraded.production_service.update(
        current.id,
        replace(valid_form(upgraded, numero_op="8101"), status=OpStatus.PRIORIDADE),
        current.row_version,
    )
    assert updated.status is OpStatus.PRIORIDADE
    assert upgraded.database.schema_version() == upgraded.database.SCHEMA_VERSION == Database.SCHEMA_VERSION
    with sqlite3.connect(db_path) as connection:
        tables = {row[0] for row in connection.execute("SELECT name FROM sqlite_master WHERE type='table'")}
    assert "op_mutations" not in tables
    assert "station_presence" not in tables
    assert list((tmp_path / "backups").glob(f"producao_operacional_pre_v{Database.SCHEMA_VERSION}_*.db"))


def test_status_values_from_legacy_database_are_normalized_before_edit(tmp_path: Path):
    container = make_container(tmp_path)
    created = container.production_service.create(valid_form(container, numero_op="8102"))
    with sqlite3.connect(container.database.db_path) as connection:
        connection.execute("UPDATE ops SET status = 'AGUARDANDO_ADM' WHERE id = ?", (created.id,))
        connection.execute(
            "UPDATE app_meta SET meta_value = '23' WHERE meta_key = 'schema_version'"
        )
        connection.commit()
    reopened = AppContainer.create(tmp_path / "settings.json")
    detail = reopened.production_service.get(created.id)
    assert detail is not None
    assert detail.status is OpStatus.AGUARDANDO


def test_write_retries_when_sqlite_is_temporarily_locked(tmp_path: Path, monkeypatch):
    from contextlib import contextmanager
    from sqlalchemy.exc import OperationalError

    container = make_container(tmp_path)
    original = container.database.write_session
    attempts = {"count": 0}

    @contextmanager
    def flaky_write_session():
        attempts["count"] += 1
        if attempts["count"] < 3:
            raise OperationalError(
                "BEGIN IMMEDIATE",
                {},
                sqlite3.OperationalError("database is locked"),
            )
        with original() as session:
            yield session

    monkeypatch.setattr(container.database, "write_session", flaky_write_session)
    created = container.production_service.create(valid_form(container, numero_op="8103"))
    assert created.numero_op == "8103"
    assert attempts["count"] == 3


def test_sqlalchemy_background_link_is_not_shown_as_the_error_message():
    from kanban_app.application.error_reporting import friendly_error_message, root_cause_from_trace

    trace = """Traceback (most recent call last):
  File \"worker.py\", line 1, in run
sqlite3.OperationalError: no such table: op_mutations

The above exception was the direct cause of the following exception:

sqlalchemy.exc.OperationalError: (sqlite3.OperationalError) no such table: op_mutations
[SQL: INSERT INTO op_mutations (...)]
(Background on this error at: https://sqlalche.me/e/20/e3q8)
"""
    assert root_cause_from_trace(trace) == "no such table: op_mutations"
    message = friendly_error_message(trace)
    assert "estrutura do banco" in message.casefold()
    assert "sqlalche.me" not in message


def test_newer_database_schema_is_not_downgraded(tmp_path: Path):
    container = make_container(tmp_path)
    with sqlite3.connect(container.database.db_path) as connection:
        connection.execute(
            "UPDATE app_meta SET meta_value = '999' WHERE meta_key = 'schema_version'"
        )
        connection.commit()
    database = Database(container.database.db_path, backups_dir=tmp_path / "backups")
    with pytest.raises(RuntimeError, match="mais nova"):
        database.create_schema()
    assert database.is_read_only()
    assert not database.read_only_is_recoverable()
