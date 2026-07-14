from __future__ import annotations

import os
import re
import shutil
import sys
import zipfile
from dataclasses import replace
from datetime import date
from pathlib import Path
from xml.etree import ElementTree

from pypdf import PdfReader

from kanban_app.application.dto import ImportPreviewDTO, OpFormDTO
from kanban_app.domain.enums import OpStatus
from kanban_app.formatting import normalize_voltage_value, parse_br_date


class DocumentImportService:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".odt"}
    REQUIRED_FIELDS: tuple[tuple[str, str], ...] = (
        ("Número da OP", "numero_op"),
        ("Cliente", "cliente"),
        ("Modelo", "modelo"),
        ("Quantidade", "quantidade"),
        ("Voltagem", "voltagem"),
        ("Prazo de entrega", "data_entrega"),
    )
    _MONTHS = {
        "JANEIRO": 1,
        "FEVEREIRO": 2,
        "MARCO": 3,
        "MARÇO": 3,
        "ABRIL": 4,
        "MAIO": 5,
        "JUNHO": 6,
        "JULHO": 7,
        "AGOSTO": 8,
        "SETEMBRO": 9,
        "OUTUBRO": 10,
        "NOVEMBRO": 11,
        "DEZEMBRO": 12,
    }

    def extract(self, source_path: str | Path, *, default_sector_id: str | None) -> ImportPreviewDTO:
        path = Path(source_path)
        if path.suffix.casefold() not in self.SUPPORTED_EXTENSIONS:
            return ImportPreviewDTO(str(path), OpFormDTO(), errors=("Formato não suportado. Use PDF, DOCX ou ODT.",))
        if not path.exists() or not path.is_file():
            return ImportPreviewDTO(str(path), OpFormDTO(), errors=("Arquivo não encontrado.",))
        try:
            raw_text = self._read_text(path)
        except Exception as exc:
            return ImportPreviewDTO(str(path), OpFormDTO(), errors=(f"Não foi possível ler o arquivo: {exc}",))

        form, missing = self._parse_text(raw_text, default_sector_id=default_sector_id)
        ocr_error = ""
        if path.suffix.casefold() == ".pdf" and self._needs_ocr(form, missing, raw_text):
            try:
                ocr_text = self._read_pdf_with_ocr(path)
            except Exception as exc:
                ocr_text = ""
                ocr_error = self._friendly_ocr_error(exc)
            if ocr_text:
                retry_form, retry_missing = self._parse_text(ocr_text, default_sector_id=default_sector_id)
                if self._quality_score(retry_form, retry_missing) > self._quality_score(form, missing):
                    form, missing = retry_form, retry_missing

        if not form.numero_op:
            filename_number = self._number_from_filename(path.stem)
            if filename_number:
                form = replace(form, numero_op=filename_number)
                missing = self.missing_fields_for_form(form)

        errors: tuple[str, ...] = ()
        if not str(raw_text or "").strip() and not form.numero_op:
            message = "O documento não contém texto legível."
            if path.suffix.casefold() == ".pdf":
                message += f" {ocr_error or 'Instale Tesseract e Poppler para importar PDFs digitalizados.'}"
            errors = (message.strip(),)
        return ImportPreviewDTO(source_path=str(path), form=form, missing_fields=tuple(missing), errors=errors)

    @classmethod
    def missing_fields_for_form(cls, form: OpFormDTO) -> list[str]:
        missing: list[str] = []
        for label, attribute in cls.REQUIRED_FIELDS:
            value = getattr(form, attribute)
            if value in (None, ""):
                missing.append(label)
        return missing

    def _read_text(self, path: Path) -> str:
        extension = path.suffix.casefold()
        if extension == ".pdf":
            reader = PdfReader(str(path))
            # A OP operacional é definida pela primeira página. Não misturar
            # anexos, desenhos ou observações das páginas seguintes.
            if not reader.pages:
                return ""
            return reader.pages[0].extract_text() or ""
        if extension == ".docx":
            from docx import Document

            document = Document(str(path))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            blocks.extend(" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows)
            return "\n".join(blocks)
        with zipfile.ZipFile(path) as archive:
            content = archive.read("content.xml")
        root = ElementTree.fromstring(content)

        # Em ODT, uma única linha visual costuma ser dividida em vários spans.
        # Iterar cada fragmento separadamente quebrava valores reais como
        # ``OP – 5273``, ``440V`` e ``08/09/2026`` em várias linhas. Agrupar
        # todo o texto de cada parágrafo preserva a estrutura visível do arquivo.
        paragraphs: list[str] = []
        for element in root.iter():
            local_name = element.tag.rsplit("}", 1)[-1]
            if local_name not in {"p", "h"}:
                continue
            paragraph = re.sub(r"\s+", " ", "".join(element.itertext())).strip()
            if paragraph and paragraph.casefold() != "objeto ole":
                paragraphs.append(paragraph)
        if paragraphs:
            return "\n".join(paragraphs)
        return "\n".join(fragment.strip() for fragment in root.itertext() if fragment.strip())

    @classmethod
    def _read_pdf_with_ocr(cls, path: Path) -> str:
        from pdf2image import convert_from_path
        import pytesseract

        tesseract = cls._find_tesseract()
        if tesseract:
            pytesseract.pytesseract.tesseract_cmd = str(tesseract)
        poppler = cls._find_poppler_bin()
        pages = convert_from_path(
            str(path),
            first_page=1,
            last_page=1,
            dpi=240,
            poppler_path=str(poppler) if poppler else None,
        )
        blocks: list[str] = []
        for image in pages:
            try:
                blocks.append(pytesseract.image_to_string(image, lang="por+eng"))
            except Exception:
                blocks.append(pytesseract.image_to_string(image))
        return "\n".join(blocks)

    @classmethod
    def _find_tesseract(cls) -> Path | None:
        candidates = [
            cls._runtime_root() / "tools" / "tesseract" / "tesseract.exe",
            Path(os.environ.get("PROGRAMFILES", "")) / "Tesseract-OCR" / "tesseract.exe",
            Path(os.environ.get("LOCALAPPDATA", "")) / "Programs" / "Tesseract-OCR" / "tesseract.exe",
        ]
        resolved = shutil.which("tesseract")
        if resolved:
            candidates.insert(0, Path(resolved))
        return next((candidate for candidate in candidates if str(candidate) and candidate.exists()), None)

    @classmethod
    def _find_poppler_bin(cls) -> Path | None:
        candidates = [
            cls._runtime_root() / "tools" / "poppler" / "Library" / "bin",
            cls._runtime_root() / "tools" / "poppler" / "bin",
        ]
        return next((candidate for candidate in candidates if candidate.exists()), None)

    @staticmethod
    def _runtime_root() -> Path:
        if getattr(sys, "frozen", False):
            return Path(sys.executable).resolve().parent
        return Path(__file__).resolve().parents[3]

    def _parse_text(self, raw_text: str, *, default_sector_id: str | None) -> tuple[OpFormDTO, list[str]]:
        text = self._normalize_text(raw_text)
        number = self._extract_op_number(text)
        client = self._line_value(text, ("CLIENTE", "RAZÃO SOCIAL", "RAZAO SOCIAL", "EMPRESA"))
        if not client:
            # DESTINO só é fallback: em vários documentos ele é cidade/UF.
            client = self._line_value(text, ("DESTINATÁRIO", "DESTINATARIO"))
        quantity = self._extract_quantity(text)
        delivery = self._extract_delivery_date(text)
        voltage = self._extract_voltage(text)
        model = self._line_value(
            text,
            ("MODELO", "EQUIPAMENTO", "PRODUTO", "DESCRIÇÃO", "DESCRICAO"),
            allow_continuation=True,
        )
        if not model:
            model = self._infer_model(text, number, client, str(quantity or ""), voltage)
        form = OpFormDTO(
            numero_op=number,
            cliente=self._clean_client(client),
            modelo=self._remove_voltage(model, voltage),
            quantidade=quantity,
            voltagem=voltage,
            data_inicio=date.today(),
            data_entrega=delivery,
            setor_id=default_sector_id,
            status=OpStatus.EM_DIA,
        )
        return form, self.missing_fields_for_form(form)

    @staticmethod
    def _normalize_text(raw_text: str) -> str:
        text = str(raw_text or "")
        replacements = {
            "–": "-",
            "—": "-",
            "−": "-",
            "º": "°",
            "\u00a0": " ",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        normalized_lines: list[str] = []
        for line in text.splitlines():
            cleaned = re.sub(r"[ \t]+", " ", line).strip()
            if cleaned:
                normalized_lines.append(cleaned)
        return "\n".join(normalized_lines)

    @staticmethod
    def _extract_op_number(text: str) -> str:
        patterns = (
            r"(?im)(?:\bO\.?\s*P\.?\b|ORDEM\s+DE\s+PRODU[ÇC][AÃ]O(?:\s+N[°O.]*)?|N[°O.]?\s*(?:DA\s+)?OP)\s*[:#-]?\s*(\d{3,})",
            r"(?im)^\s*OP\s*[-:]\s*(\d{3,})\s*$",
        )
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        return ""

    @classmethod
    def _line_value(
        cls,
        text: str,
        labels: tuple[str, ...],
        *,
        allow_continuation: bool = False,
    ) -> str:
        lines = text.splitlines()
        label_expression = "|".join(re.escape(label) for label in labels)
        known_labels = (
            "OP", "CLIENTE", "RAZÃO SOCIAL", "RAZAO SOCIAL", "EMPRESA",
            "DESTINO", "DESTINATÁRIO", "DESTINATARIO", "MODELO", "EQUIPAMENTO",
            "PRODUTO", "DESCRIÇÃO", "DESCRICAO", "QUANTIDADE", "QTD", "QTDE",
            "VOLTAGEM", "TENSÃO", "TENSAO", "PRAZO", "ENTREGA", "DATA DE ENTREGA",
            "PESO", "CÓDIGO", "CODIGO", "PEDIDO", "NOTA", "CERTIFICADO", "OBS",
            "VALOR", "TOTAL", "VOLUME", "PALLET", "MOTOR N.S", "PAINEL N.S",
        )
        next_label_expression = "|".join(re.escape(label) for label in known_labels)

        for index, line in enumerate(lines):
            # Aceita CLIENTE:, CLIENTE - e células convertidas para "CLIENTE | valor".
            match = re.match(rf"(?i)^\s*(?:{label_expression})\s*(?:[:#-]|\|)?\s*(.*?)\s*$", line)
            if not match:
                continue
            parts: list[str] = []
            value = cls._clean_field_value(match.group(1))
            if value:
                parts.append(value)
            if allow_continuation:
                # Alguns modelos reais quebram a descrição em duas linhas. Junta
                # somente até o próximo campo reconhecido para não capturar o
                # restante do documento.
                for candidate_line in lines[index + 1 : index + 4]:
                    candidate = cls._clean_field_value(candidate_line)
                    if not candidate:
                        continue
                    if re.match(rf"(?i)^\s*(?:{next_label_expression})\b\s*(?:[:#-]|\|)?", candidate):
                        break
                    if candidate.casefold() == "objeto ole":
                        continue
                    parts.append(candidate)
            elif not parts and index + 1 < len(lines):
                candidate = cls._clean_field_value(lines[index + 1])
                if candidate and not re.match(r"^[A-ZÁÉÍÓÚÂÊÔÃÕÇ /]{2,35}\s*[:#|-]", candidate):
                    parts.append(candidate)
            if parts:
                return cls._clean_field_value(" ".join(parts))

        # Fallback para tabelas que chegaram como uma linha longa com pipes.
        match = re.search(rf"(?i)(?:{label_expression})\s*(?:[:#-]|\|)\s*([^|\n]+)", text)
        return cls._clean_field_value(match.group(1)) if match else ""

    @staticmethod
    def _clean_field_value(value: str) -> str:
        cleaned = re.sub(r"\s+", " ", str(value or "")).strip(" -:;|")
        # Para quando a extração cola o próximo rótulo na mesma linha.
        cleaned = re.split(
            r"\s+(?=(?:CLIENTE|DESTINO|MODELO|EQUIPAMENTO|QUANTIDADE|QTD|VOLTAGEM|PRAZO|ENTREGA|C[ÓO]DIGO|OBS\.?)[ :#|-])",
            cleaned,
            maxsplit=1,
            flags=re.IGNORECASE,
        )[0]
        return cleaned.strip(" -:;|")

    @staticmethod
    def _clean_client(value: str) -> str:
        cleaned = re.sub(r"\s+", " ", str(value or "")).strip(" -:;|")
        return cleaned.rstrip(".")

    @staticmethod
    def _extract_quantity(text: str) -> int | None:
        patterns = (
            r"(?im)^\s*(?:QUANTIDADE|QTD\.?|QTDE\.?)\s*[:#|-]?\s*(\d{1,7})\b",
            r"(?i)\b(?:QUANTIDADE|QTD\.?|QTDE\.?)\s*[:#|-]?\s*(\d{1,7})\s*(?:P[ÇC]S?\.?|UN(?:IDADES?)?\.?|PCS?)?\b",
        )
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                try:
                    value = int(match.group(1))
                    return value if value > 0 else None
                except ValueError:
                    pass
        return None

    @classmethod
    def _extract_delivery_date(cls, text: str) -> date | None:
        labels = r"PRAZO\s+DE\s+ENTREGA|DATA\s+DE\s+ENTREGA|PREVIS[ÃA]O\s+DE\s+ENTREGA|ENTREGA|PRAZO"
        numeric = re.search(
            rf"(?im)(?:{labels})[^\d\n]{{0,30}}(\d{{8}}|\d{{1,2}}[./-]\d{{1,2}}(?:[./-]\d{{2,4}})?)",
            text,
        )
        if numeric:
            parsed = cls._parse_date(numeric.group(1))
            if parsed:
                return parsed
        words = re.search(
            rf"(?im)(?:{labels})[^\d\n]{{0,30}}(\d{{1,2}})\s+DE\s+([A-ZÇÃÕÉÊÍÓÔÚ]+)\s+DE\s+(\d{{4}})",
            text,
        )
        if words:
            month = cls._MONTHS.get(words.group(2).upper())
            if month:
                try:
                    return date(int(words.group(3)), month, int(words.group(1)))
                except ValueError:
                    pass
        return None

    @staticmethod
    def _extract_voltage(text: str) -> str:
        patterns = (
            r"(?im)^\s*(?:VOLTAGEM|TENS[ÃA]O)\s*[:#|-]?\s*((?:N\s*/?\s*A)|(?:NÃO|NAO)\s+APLICÁVEL|\d{2,4}(?:\s*[/\\-]\s*\d{2,4})?\s*(?:V|VOLTS?|VAC)?)\b",
            r"\b(\d{2,4}(?:\s*[/\\-]\s*\d{2,4})?\s*(?:V|VOLTS?|VAC))\b",
        )
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = normalize_voltage_value(match.group(1))
                if value:
                    return value
        return ""

    @staticmethod
    def _remove_voltage(model: str, voltage: str) -> str:
        cleaned = re.sub(r"\b\d{2,4}(?:\s*[/\\-]\s*\d{2,4})?\s*(?:V|VOLTS?|VAC)\b", "", str(model or ""), flags=re.IGNORECASE)
        cleaned = re.sub(r"\s+", " ", cleaned).strip(" -|,;")
        return cleaned

    @staticmethod
    def _infer_model(text: str, *known: str) -> str:
        ignored = {str(value or "").strip() for value in known if value}
        label_words = (
            "ORDEM DE PRODUÇÃO",
            "ORDEM DE PRODUCAO",
            "NÚMERO DA OP",
            "NUMERO DA OP",
            "CLIENTE",
            "DESTINO",
            "QUANTIDADE",
            "VOLTAGEM",
            "TENSÃO",
            "PRAZO",
            "ENTREGA",
            "CÓDIGO",
            "CODIGO",
            "CERTIFICADO",
            "OBS",
            "NOTA",
            "PESO",
            "TOTAL",
            "RIO DE JANEIRO",
        )
        candidates: list[str] = []
        for line in text.splitlines():
            cleaned = re.sub(r"\s+", " ", line).strip(" -|,;")
            upper = cleaned.upper()
            if not (8 <= len(cleaned) <= 240):
                continue
            if any(word in upper for word in label_words):
                continue
            if any(value and value.casefold() in cleaned.casefold() for value in ignored):
                continue
            if re.fullmatch(r"[\d\s,./-]+", cleaned):
                continue
            candidates.append(cleaned)
        # Modelos normalmente combinam letras e números; prioriza esse padrão.
        for candidate in candidates:
            if re.search(r"[A-Za-zÁ-ÿ]", candidate) and re.search(r"\d", candidate):
                return candidate
        return candidates[0] if candidates else ""

    @staticmethod
    def _number_from_filename(stem: str) -> str:
        match = re.search(r"(?i)\b(?:O\.?P\.?)\s*[-_ ]*([0-9]{3,})\b", stem)
        return match.group(1) if match else ""

    @staticmethod
    def _parse_date(value: str) -> date | None:
        return parse_br_date(value)

    @staticmethod
    def _needs_ocr(form: OpFormDTO, missing: list[str], raw_text: str) -> bool:
        return len(str(raw_text or "").strip()) < 40 or not form.numero_op or len(missing) >= 3

    @staticmethod
    def _quality_score(form: OpFormDTO, missing: list[str]) -> int:
        score = 100 - len(missing) * 15
        score += 10 if form.numero_op else 0
        score += min(20, len(form.modelo or "") // 8)
        score += min(10, len(form.cliente or "") // 8)
        return score

    @staticmethod
    def _friendly_ocr_error(exc: Exception) -> str:
        message = str(exc or "").casefold()
        if "tesseract" in message:
            return "Tesseract não foi encontrado para ler o PDF digitalizado."
        if "poppler" in message or "page count" in message:
            return "Poppler não foi encontrado para converter o PDF digitalizado."
        return "Não foi possível executar o OCR deste PDF."
